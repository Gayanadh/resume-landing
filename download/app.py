import os, io, json, tempfile, subprocess
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz
from groq import Groq

app = Flask(__name__)
CORS(app)

groq_client = Groq(api_key=os.environ.get("GROQ_API_KEY")) if os.environ.get("GROQ_API_KEY") else None

ATS_ORDER = [
    "Professional Summary","Objective","Key Strengths","Skills","Technical Skills",
    "Achievements","Certifications","Professional History","Work Experience",
    "Education","Projects","Awards","Languages","Volunteer"
]

# ─────────────────────────────────────────────────────────────
#  TEXT EXTRACTION
# ─────────────────────────────────────────────────────────────
def extract_text_pdf(file_bytes):
    pdf = fitz.open(stream=file_bytes, filetype="pdf")
    return "\n".join(p.get_text() for p in pdf)

def extract_text_docx(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines)

# ─────────────────────────────────────────────────────────────
#  GROQ — classify + polish
# ─────────────────────────────────────────────────────────────
def classify_and_polish(resume_text, new_text, template):
    existing = []
    for line in resume_text.split("\n"):
        l = line.strip()
        if l and len(l) < 40 and not l.startswith("•") and not l[0].isdigit():
            for h in ATS_ORDER:
                if h.lower() in l.lower():
                    existing.append(l)
                    break

    style = {
        "google":  "XYZ formula: Accomplished [X] as measured by [Y] by doing [Z].",
        "apple":   "Elegant, product-centric, user-impact focus.",
        "amazon":  "High-scale metrics and Amazon Leadership Principles."
    }.get(template, "Strong action verb + quantified result.")

    r = groq_client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role":"system","content":f"""You are an ATS resume expert.
Resume headings found: {json.dumps(existing if existing else ATS_ORDER)}
Choose the most appropriate heading for this sentence. Prefer existing headings from the resume.
Rewrite the sentence: {style}
Return ONLY JSON: {{"heading":"...","polished":"..."}}"""},
        {"role":"user","content":new_text}],
        max_tokens=200, temperature=0.3
    )
    try:
        return json.loads(r.choices[0].message.content.strip())
    except:
        return {"heading":"Achievements","polished":new_text}

# ─────────────────────────────────────────────────────────────
#  DOCX INJECTION — finds heading via XML, inserts after it
# ─────────────────────────────────────────────────────────────
def inject_docx(file_bytes, target_heading, new_text):
    doc = Document(io.BytesIO(file_bytes))

    # Find the best paragraph to insert after
    insert_after_idx = None
    last_content_idx = None

    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        if target_heading.lower() in txt:
            insert_after_idx = i
        if insert_after_idx is not None and i > insert_after_idx:
            # If we've found the heading, track last bullet under it
            if any(h.lower() in txt for h in ATS_ORDER if h.lower() != target_heading.lower()):
                break  # hit next heading
            if txt:
                last_content_idx = i

    # Insert after last bullet, or right after heading
    anchor_idx = last_content_idx if last_content_idx is not None else insert_after_idx

    if anchor_idx is not None:
        anchor = doc.paragraphs[anchor_idx]
    else:
        # Heading doesn't exist — find position in ATS order, create it
        ats_i = ATS_ORDER.index(target_heading) if target_heading in ATS_ORDER else 99
        anchor_idx = len(doc.paragraphs) - 1
        for i, p in enumerate(doc.paragraphs):
            for nh in ATS_ORDER[ats_i + 1:]:
                if nh.lower() in p.text.strip().lower():
                    anchor_idx = i
                    break
            if anchor_idx != len(doc.paragraphs) - 1:
                break

        anchor = doc.paragraphs[anchor_idx]

        # Build heading XML
        h_el = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "FFFFFF")
        pPr.append(shd)
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "120")
        sp.set(qn("w:after"), "60")
        pPr.append(sp)
        h_el.append(pPr)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        col = OxmlElement("w:color")
        col.set(qn("w:val"), "2E74B5")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        rPr.extend([b, col, sz])
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = target_heading + ":"
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        run.append(t)
        h_el.append(run)
        anchor._element.addprevious(h_el)
        anchor = doc.paragraphs[anchor_idx]

    # Build bullet XML and insert after anchor
    b_el = OxmlElement("w:p")
    pPr2 = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    pPr2.append(ind)
    sp2 = OxmlElement("w:spacing")
    sp2.set(qn("w:before"), "40")
    sp2.set(qn("w:after"), "40")
    pPr2.append(sp2)
    b_el.append(pPr2)
    run2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    sz2 = OxmlElement("w:sz")
    sz2.set(qn("w:val"), "20")
    rPr2.append(sz2)
    run2.append(rPr2)
    t2 = OxmlElement("w:t")
    t2.text = f"\u2022  {new_text}"
    t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run2.append(t2)
    b_el.append(run2)
    anchor._element.addnext(b_el)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ─────────────────────────────────────────────────────────────
#  PDF INJECTION — convert to DOCX, inject, convert back
#  This is the ONLY reliable way to avoid overlapping content
# ─────────────────────────────────────────────────────────────
def inject_pdf(file_bytes, target_heading, new_text):
    with tempfile.TemporaryDirectory() as d:
        # Step 1: PDF → DOCX via LibreOffice
        pdf_path = os.path.join(d, "input.pdf")
        open(pdf_path, "wb").write(file_bytes)
        subprocess.run(
            ["libreoffice","--headless","--convert-to","docx","--outdir",d,pdf_path],
            capture_output=True, timeout=60
        )
        docx_path = os.path.join(d, "input.docx")
        if not os.path.exists(docx_path):
            # LibreOffice not available — return original unchanged
            return io.BytesIO(file_bytes)

        # Step 2: Inject into DOCX
        docx_bytes = open(docx_path,"rb").read()
        enhanced   = inject_docx(docx_bytes, target_heading, new_text)

        # Step 3: DOCX → PDF
        out_docx = os.path.join(d, "enhanced.docx")
        open(out_docx, "wb").write(enhanced.read())
        subprocess.run(
            ["libreoffice","--headless","--convert-to","pdf","--outdir",d,out_docx],
            capture_output=True, timeout=60
        )
        out_pdf = os.path.join(d, "enhanced.pdf")
        if os.path.exists(out_pdf):
            result = io.BytesIO(open(out_pdf,"rb").read())
            result.seek(0)
            return result

    return io.BytesIO(file_bytes)

# ─────────────────────────────────────────────────────────────
#  COMPLETE AI REBUILD — Groq reads entire resume and rewrites
# ─────────────────────────────────────────────────────────────
def rebuild_resume(resume_text, template, job_desc):
    style = {
        "google":  "XYZ formula for all achievements. Minimalist, data-driven tone.",
        "apple":   "Elegant product-centric language. Focus on user impact.",
        "amazon":  "Amazon Leadership Principles. Scale, metrics, ownership."
    }.get(template, "Strong action verbs, quantified results.")

    jd_note = f"\n\nOptimize keywords for this job description:\n{job_desc[:1000]}" if job_desc else ""

    r = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"system","content":f"""You are a world-class ATS resume writer.
Rewrite this entire resume to be ATS-optimized, professional and impactful.
Style: {style}
Rules:
- Keep ALL original facts, companies, dates, names — never invent
- Use strong action verbs for every bullet
- Add metrics/numbers where implied
- Improve clarity and conciseness
- Keep all sections and headings from original
- Return ONLY valid JSON, no other text:
{{
  "name": "...",
  "contact": {{"email":"...","phone":"...","location":"..."}},
  "sections": [
    {{"heading":"...","bullets":["...","..."]}},
    {{"heading":"Professional History","is_jobs":true,"jobs":[
      {{"title":"...","company":"...","period":"...","bullets":["..."]}}
    ]}}
  ]
}}{jd_note}"""},
        {"role":"user","content":f"Resume to rewrite:\n{resume_text}"}],
        max_tokens=4000, temperature=0.5
    )
    raw = r.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"): raw=raw[4:]
    return json.loads(raw.strip())

def build_docx_from_json(data):
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.75)

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    nr = name_p.add_run(data.get("name",""))
    nr.bold = True; nr.font.size = Pt(20); nr.font.color.rgb = RGBColor(26,58,107)

    # Contact
    c = data.get("contact",{})
    parts = [v for v in c.values() if v]
    if parts:
        cp = doc.add_paragraph("  •  ".join(parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        for r in cp.runs:
            r.font.size = Pt(9); r.font.color.rgb = RGBColor(100,100,100)

    def section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text.upper())
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(46,116,181)
        from docx.oxml import OxmlElement
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'2E74B5')
        pBdr.append(bot); pPr.append(pBdr)

    def bullet(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.15)
        r = p.add_run(f"• {text}")
        r.font.size = Pt(10)

    for sec in data.get("sections",[]):
        section_heading(sec.get("heading",""))
        if sec.get("is_jobs"):
            for job in sec.get("jobs",[]):
                jp = doc.add_paragraph()
                jp.paragraph_format.space_before = Pt(4)
                jr = jp.add_run(f"{job.get('title','')}  —  {job.get('company','')}")
                jr.bold = True; jr.font.size = Pt(10)
                per = jp.add_run(f"   {job.get('period','')}")
                per.font.size = Pt(9); per.font.color.rgb = RGBColor(130,130,130)
                for b in job.get("bullets",[]): bullet(b)
        else:
            for b in sec.get("bullets",[]): bullet(b)

    out = io.BytesIO(); doc.save(out); out.seek(0)
    return out

# ─────────────────────────────────────────────────────────────
#  ROUTES
# ─────────────────────────────────────────────────────────────
@app.route("/parse", methods=["POST"])
def parse_resume():
    """Extract and parse resume content, return structured JSON for the frontend editor."""
    try:
        file = request.files.get("resume")
        if not file:
            return jsonify({"error": "No file uploaded"}), 400

        file_bytes = file.read()
        fn = file.filename.lower()

        # Extract raw text
        if fn.endswith(".pdf"):
            resume_text = extract_text_pdf(file_bytes)
        elif fn.endswith(".docx"):
            resume_text = extract_text_docx(file_bytes)
        else:
            return jsonify({"error": "Unsupported file type. Please upload .pdf or .docx"}), 400

        if not resume_text.strip():
            return jsonify({"error": "Could not extract text from the file. It may be image-based."}), 400

        # Use Groq to parse into structured JSON
        if groq_client:
            r = groq_client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[{
                    "role": "system",
                    "content": """You are a resume parser. Extract ALL information from this resume text and return it as structured JSON.
Return ONLY valid JSON with this exact structure:
{
  "name": "Full Name",
  "title": "Job Title or current role",
  "email": "email@example.com",
  "phone": "+1 234 567 890",
  "location": "City, Country",
  "linkedin": "linkedin URL or empty string",
  "website": "website URL or empty string",
  "summary": "The professional summary or objective text",
  "skills": [{"name": "Skill Name", "level": 85}],
  "experience": [{"title": "Job Title", "company": "Company", "period": "2020 - Present", "bullets": ["achievement 1", "achievement 2"]}],
  "education": [{"degree": "Degree Name", "school": "School Name", "detail": "2015 - 2019, Honors"}],
  "tools": ["Tool1", "Tool2"],
  "languages": [{"name": "English", "level": 5}, {"name": "Spanish", "level": 3}]
}
Rules:
- Extract EVERY piece of information you find
- For skills level, estimate 50-95 based on how prominently they're featured
- For language level, use 1-5 scale (5=native)
- Keep all bullet points exactly as written
- If a field is not found, use empty string for strings, empty array for arrays
- Return ONLY the JSON, no other text"""
                }, {
                    "role": "user",
                    "content": f"Parse this resume:\n{resume_text}"
                }],
                max_tokens=3000, temperature=0.1
            )
            raw = r.choices[0].message.content.strip()
            if raw.startswith("```"):
                raw = raw.split("```")[1]
                if raw.startswith("json"):
                    raw = raw[4:]
            try:
                data = json.loads(raw.strip())
            except:
                data = {"name": "", "title": "", "email": "", "phone": "", "location": "",
                        "linkedin": "", "website": "", "summary": resume_text[:500],
                        "skills": [], "experience": [], "education": [], "tools": [], "languages": []}
        else:
            # No Groq - return raw text
            data = {"name": "", "title": "", "email": "", "phone": "", "location": "",
                    "linkedin": "", "website": "", "summary": resume_text[:500],
                    "skills": [], "experience": [], "education": [], "tools": [], "languages": []}

        data["raw_text"] = resume_text
        data["filename"] = file.filename
        return jsonify(data)

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/polish", methods=["POST"])
def polish():
    if not groq_client:
        return jsonify({"polished_text":"GROQ_API_KEY not set","heading":"Achievements"}),500
    try:
        file     = request.files.get("resume")
        text     = request.form.get("text","")
        template = request.form.get("template","google")
        if not file or not text:
            return jsonify({"polished_text":"Upload resume and enter text","heading":"Achievements"}),400

        file_bytes = file.read()
        fn = file.filename.lower()
        resume_text = extract_text_pdf(file_bytes) if fn.endswith(".pdf") else extract_text_docx(file_bytes)
        result = classify_and_polish(resume_text, text, template)
        return jsonify({"polished_text": result["polished"], "heading": result["heading"]})
    except Exception as e:
        return jsonify({"polished_text":f"Error: {e}","heading":"Achievements"}),200

@app.route("/upgrade", methods=["POST"])
def upgrade():
    try:
        file     = request.files.get("resume")
        updates  = request.form.get("updates","").strip()
        heading  = request.form.get("heading","Achievements").strip()
        fmt      = request.form.get("format","docx").lower()
        template = request.form.get("template","google")
        if not file: return jsonify({"error":"No file"}),400

        file_bytes = file.read()
        fn = file.filename.lower()

        # If no heading yet, classify now
        if not heading and groq_client:
            rt = extract_text_pdf(file_bytes) if fn.endswith(".pdf") else extract_text_docx(file_bytes)
            result  = classify_and_polish(rt, updates, template)
            heading = result["heading"]
            updates = result["polished"]

        if fn.endswith(".pdf"):
            # Always go PDF→DOCX→inject→PDF — no more drawing on top
            result = inject_pdf(file_bytes, heading, updates)
            return send_file(result, mimetype="application/pdf",
                             as_attachment=True, download_name="Optimized_Resume.pdf")

        elif fn.endswith(".docx"):
            enhanced = inject_docx(file_bytes, heading, updates)
            if fmt == "docx":
                return send_file(enhanced,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    as_attachment=True, download_name="Optimized_Resume.docx")
            # DOCX→PDF
            enhanced.seek(0)
            with tempfile.TemporaryDirectory() as d:
                src = os.path.join(d,"r.docx")
                open(src,"wb").write(enhanced.read())
                subprocess.run(["libreoffice","--headless","--convert-to","pdf","--outdir",d,src],
                               capture_output=True, timeout=60)
                pp = os.path.join(d,"r.pdf")
                if os.path.exists(pp):
                    return send_file(io.BytesIO(open(pp,"rb").read()),
                        mimetype="application/pdf", as_attachment=True, download_name="Optimized_Resume.pdf")
            enhanced.seek(0)
            return send_file(enhanced,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True, download_name="Optimized_Resume.docx")

        return jsonify({"error":"Unsupported file type"}),400
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error":str(e)}),500

@app.route("/rebuild", methods=["POST"])
def rebuild():
    if not groq_client:
        return jsonify({"error":"GROQ_API_KEY not set"}),500
    try:
        file     = request.files.get("resume")
        template = request.form.get("template","google")
        job_desc = request.form.get("job_desc","")
        fmt      = request.form.get("format","docx").lower()
        if not file: return jsonify({"error":"No file"}),400

        file_bytes = file.read()
        fn = file.filename.lower()
        resume_text = extract_text_pdf(file_bytes) if fn.endswith(".pdf") else extract_text_docx(file_bytes)

        data    = rebuild_resume(resume_text, template, job_desc)
        docx_io = build_docx_from_json(data)

        if fmt == "pdf":
            with tempfile.TemporaryDirectory() as d:
                src = os.path.join(d,"rebuilt.docx")
                docx_io.seek(0); open(src,"wb").write(docx_io.read())
                subprocess.run(["libreoffice","--headless","--convert-to","pdf","--outdir",d,src],
                               capture_output=True, timeout=60)
                pp = os.path.join(d,"rebuilt.pdf")
                if os.path.exists(pp):
                    return send_file(io.BytesIO(open(pp,"rb").read()),
                        mimetype="application/pdf", as_attachment=True, download_name="Rebuilt_Resume.pdf")
            docx_io.seek(0)

        docx_io.seek(0)
        return send_file(docx_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True, download_name="Rebuilt_Resume.docx")

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error":str(e)}),500

if __name__ == "__main__":
    port = int(os.environ.get("PORT",10000))
    app.run(host="0.0.0.0", port=port, debug=False)
